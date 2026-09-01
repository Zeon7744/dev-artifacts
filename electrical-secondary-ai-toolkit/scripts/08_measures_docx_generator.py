#!/usr/bin/env python3
"""
二次措施单生成器 - 带Word文档导出功能
基于知识图谱和模板匹配，生成标准格式的Word文档
"""

import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from datetime import datetime
import copy

try:
    from docx import Document
    from docx.shared import Pt, Inches, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，Word导出功能不可用")


class MeasuresDocGenerator:
    """二次措施单Word文档生成器"""
    
    # 措施单表格标准格式（9列）
    TABLE_HEADERS = [
        '序号', '执行', '时间', 
        '安全技术措施内容', '安全技术措施内容', 
        '安全技术措施内容', '安全技术措施内容',
        '恢复', '时间'
    ]
    
    def __init__(self):
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx: pip install python-docx")
    
    def generate(self, measures_data: Dict[str, Any], 
                 output_path: str,
                 add_signatures: bool = True) -> str:
        """生成Word文档"""
        
        # 创建文档
        doc = Document()
        
        # 设置文档标题
        title = doc.add_heading(f"{measures_data.get('station_name', 'XX变电站')} 厂站二次设备及回路工作安全技术措施单", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加免责声明
        disclaimer = doc.add_paragraph()
        disclaimer_text = "（本措施单仅用于加深正文条款的认识，不作为现场安全措施的制定依据，现场应根据图纸及实际情况制定安全措施）"
        run = disclaimer.add_run(disclaimer_text)
        run.font.size = Pt(9)
        run.font.color.rgb = None  # 灰色
        disclaimer.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加基本信息表格
        self._add_basic_info(doc, measures_data)
        
        # 添加措施单主体表格
        self._add_measures_table(doc, measures_data.get('items', []))
        
        # 添加签名栏
        if add_signatures:
            self._add_signatures(doc, measures_data)
        
        # 保存文档
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        doc.save(str(output_path))
        
        return str(output_path)
    
    def _add_basic_info(self, doc: Document, data: Dict[str, Any]):
        """添加基本信息"""
        info_table = doc.add_table(rows=3, cols=4)
        info_table.style = 'Table Grid'
        
        # 第一行：措施单编号、工作票编号
        info_table.rows[0].cells[0].text = '措施单编号：'
        info_table.rows[0].cells[1].text = data.get('sheet_number', '')
        info_table.rows[0].cells[2].text = '工作票编号：'
        info_table.rows[0].cells[3].text = data.get('work_ticket_number', '')
        
        # 第二行：编制人、审核人
        info_table.rows[1].cells[0].text = '编制人：'
        info_table.rows[1].cells[1].text = data.get('preparer', '')
        info_table.rows[1].cells[2].text = '审核人：'
        info_table.rows[1].cells[3].text = data.get('approver', '')
        
        # 第三行：备注
        info_table.rows[2].cells[0].text = '备注：'
        info_table.rows[2].cells[1].colspan = 3
        info_table.rows[2].cells[1].text = data.get('notes', '')
        
        # 设置列宽
        for row in info_table.rows:
            row.cells[0].width = Cm(3)
            row.cells[1].width = Cm(4)
            row.cells[2].width = Cm(3)
            row.cells[3].width = Cm(6)
        
        doc.add_paragraph()  # 空行
    
    def _add_measures_table(self, doc: Document, items: List[Dict]):
        """添加措施单主体表格"""
        
        # 创建表格
        table = doc.add_table(rows=1, cols=9)
        table.style = 'Table Grid'
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        
        # 设置表头
        header_cells = table.rows[0].cells
        for i, header in enumerate(self.TABLE_HEADERS):
            header_cells[i].text = header
            # 设置表头样式
            for paragraph in header_cells[i].paragraphs:
                for run in paragraph.runs:
                    run.bold = True
                    run.font.size = Pt(9)
        
        # 添加数据行
        current_category = None
        for item in items:
            row_cells = table.add_row().cells
            
            # 序号
            row_cells[0].text = item.get('seq', '')
            
            # 执行标记（空，待填写）
            row_cells[1].text = ''
            
            # 时间（空，待填写）
            row_cells[2].text = ''
            
            # 措施内容（合并4列）
            content = item.get('content', '')
            if item.get('category') and '安措' in item.get('category', ''):
                # 类别标题行，合并后三列
                row_cells[3].text = f"[{item['category']}]"
                row_cells[3].merge(row_cells[6])
                row_cells[7].text = ''
                row_cells[8].text = ''
            else:
                row_cells[3].text = content
                # 确保内容列有足够宽度
                for cell in [row_cells[3], row_cells[4], row_cells[5], row_cells[6]]:
                    cell.width = Cm(4)
            
            # 恢复标记（空，待填写）
            row_cells[7].text = ''
            
            # 恢复时间（空，待填写）
            row_cells[8].text = ''
            
            # 设置行高
            for cell in row_cells:
                for paragraph in cell.paragraphs:
                    paragraph.paragraph_format.line_spacing = 1.2
        
        # 添加说明文字
        doc.add_paragraph()
        explanation = doc.add_paragraph()
        explanation_text = '说明：安全技术措施应按照工作顺序填写。已执行，在执行栏打"√"，已恢复，在恢复栏打"√"，并在对应的时间栏填写执行和恢复的具体时间，不需恢复的，在恢复栏打"○"，在时间栏划横杠"—"。'
        run = explanation.add_run(explanation_text)
        run.font.size = Pt(9)
        run.font.italic = True
    
    def _add_signatures(self, doc: Document, data: Dict[str, Any]):
        """添加签名栏"""
        doc.add_paragraph()  # 空行
        
        # 签名表格
        sig_table = doc.add_table(rows=2, cols=6)
        sig_table.style = 'Table Grid'
        
        # 第一行：工作负责人、执行人、监护人
        sig_labels = ['工作负责人（审批人）', '执行人', '监护人']
        sig_values = [data.get('approver', ''), data.get('executor', ''), data.get('supervisor', '')]
        
        for i, (label, value) in enumerate(zip(sig_labels, sig_values)):
            cell = sig_table.rows[0].cells[i]
            cell.text = f"{label}：{value}"
        
        # 第二行：恢复人、监护人、见证人
        sig_labels2 = ['恢复人', '监护人', '见证人']
        sig_values2 = [data.get('restorer', ''), data.get('supervisor', ''), data.get('witness', '')]
        
        for i, (label, value) in enumerate(zip(sig_labels2, sig_values2)):
            cell = sig_table.rows[1].cells[i]
            cell.text = f"{label}：{value}"
        
        # 设置列宽
        for row in sig_table.rows:
            for cell in row.cells:
                cell.width = Cm(4)


class BatchMeasuresGenerator:
    """批量措施单生成器"""
    
    def __init__(self):
        self.doc_generator = MeasuresDocGenerator()
    
    def generate_batch(self, batch_params: List[Dict[str, Any]], 
                       output_dir: str) -> List[str]:
        """批量生成措施单"""
        output_paths = []
        
        for idx, params in enumerate(batch_params):
            # 生成输出文件名
            station = params.get('substation_name', f'station_{idx}')
            equipment = params.get('equipment_type', 'unknown')
            work_type = params.get('work_type', 'unknown')
            
            filename = f"措施单_{station}_{equipment}_{work_type}_{idx+1}.docx"
            output_path = Path(output_dir) / filename
            
            # 生成措施单
            try:
                # 这里应该调用MeasuresGenerator生成measures_data
                # 为了简化，使用示例数据
                sample_data = {
                    'station_name': station,
                    'sheet_number': f"2026-{idx+1:04d}",
                    'work_ticket_number': params.get('work_ticket_number', ''),
                    'items': self._generate_sample_items(params),
                    'preparer': params.get('preparer', ''),
                    'approver': params.get('approver', ''),
                    'executor': params.get('executor', ''),
                    'supervisor': params.get('supervisor', ''),
                    'restorer': params.get('restorer', ''),
                    'witness': params.get('witness', ''),
                    'notes': params.get('notes', '')
                }
                
                path = self.doc_generator.generate(sample_data, str(output_path))
                output_paths.append(path)
                
            except Exception as e:
                print(f"生成失败: {filename}, 错误: {e}")
        
        return output_paths
    
    def _generate_sample_items(self, params: Dict) -> List[Dict]:
        """生成示例措施项（实际应用中应调用MeasuresGenerator）"""
        return [
            {'seq': '一', 'category': '', 'subcategory': '', 'content': '核实二次措施单与现场实际接线一致', 'location': params.get('substation_name', '')},
            {'seq': '', 'category': '密封安措', 'subcategory': '', 'content': '', 'location': ''},
            {'seq': '1', 'category': '密封安措', 'subcategory': '压板密封', 'content': f'逐一密封工作票中列明退出的涉及运行设备的压板', 'location': f'{params.get("substation_name", "XX")}A屏'},
            {'seq': '2', 'category': '密封安措', 'subcategory': '空开密封', 'content': '逐一密封工作票中列明工作范围内的交流电压、直流电源空开', 'location': f'{params.get("substation_name", "XX")}A屏'},
            {'seq': '', 'category': '回路安措', 'subcategory': '', 'content': '', 'location': ''},
            {'seq': '3', 'category': '回路安措', 'subcategory': '电流回路', 'content': f'确认{params.get("equipment_type", "设备")}CT无流后，打开二次电流回路端子连接片', 'location': f'{params.get("substation_name", "XX")}A屏'},
        ]


def main():
    """主函数 - 命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='二次措施单Word文档生成器')
    parser.add_argument('--input', '-i', required=True, help='输入参数JSON路径')
    parser.add_argument('--output', '-o', default='output/measures.docx', help='输出Word路径')
    parser.add_argument('--batch', '-b', action='store_true', help='批量生成模式')
    parser.add_argument('--batch-dir', '-d', default='output/batch', help='批量输出目录')
    
    args = parser.parse_args()
    
    # 读取输入
    with open(args.input, 'r', encoding='utf-8') as f:
        if args.batch:
            params_list = json.load(f)
        else:
            params_list = [json.load(f)]
    
    if args.batch:
        # 批量生成
        generator = BatchMeasuresGenerator()
        paths = generator.generate_batch(params_list, args.batch_dir)
        print(f"✓ 批量生成完成，共生成 {len(paths)} 个文档")
        for p in paths:
            print(f"  - {p}")
    else:
        # 单个生成
        params = params_list[0]
        
        # 先调用MeasuresGenerator生成数据
        from _05_measures_generation import MeasuresGenerator
        gen = MeasuresGenerator()
        measures = gen.generate(params)
        
        # 再用Word生成器导出
        doc_gen = MeasuresDocGenerator()
        output_path = doc_gen.generate(measures.to_dict(), args.output)
        
        print(f"✓ 措施单已生成: {output_path}")
        
        # 同时保存JSON
        json_path = args.output.replace('.docx', '.json')
        with open(json_path, 'w', encoding='utf-8') as f:
            json.dump(measures.to_dict(), f, ensure_ascii=False, indent=2)
        print(f"✓ JSON已保存: {json_path}")


if __name__ == '__main__':
    main()
