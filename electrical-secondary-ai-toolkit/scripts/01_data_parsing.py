#!/usr/bin/env python3
"""
电气二次标准文档解析器
用于解析广东电网二次作业标准文档，提取结构化内容
"""

import re
import json
from pathlib import Path
from typing import List, Dict, Any, Optional
from dataclasses import dataclass, asdict
from datetime import datetime

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("警告: python-docx 未安装，部分功能不可用")


@dataclass
class DocumentSection:
    """文档章节数据结构"""
    section_id: str
    title: str
    content: str
    level: int
    children: List['DocumentSection'] = None
    
    def __post_init__(self):
        if self.children is None:
            self.children = []
    
    def to_dict(self) -> dict:
        return {
            'section_id': self.section_id,
            'title': self.title,
            'content': self.content,
            'level': self.level,
            'children': [c.to_dict() for c in self.children]
        }


@dataclass
class TableData:
    """表格数据结构"""
    table_id: str
    headers: List[str]
    rows: List[List[str]]
    source_page: Optional[int] = None
    
    def to_dict(self) -> dict:
        return {
            'table_id': self.table_id,
            'headers': self.headers,
            'rows': self.rows,
            'source_page': self.source_page
        }


@dataclass 
class KnowledgeElement:
    """知识元素数据结构"""
    element_id: str
    element_type: str  # term, rule, procedure, requirement
    name: str
    definition: str
    context: str
    references: List[str]
    metadata: Dict[str, Any]
    
    def to_dict(self) -> dict:
        return asdict(self)


class SecondaryDocParser:
    """二次作业标准文档解析器"""
    
    def __init__(self, config: Optional[Dict] = None):
        self.config = config or {}
        self.sections: List[DocumentSection] = []
        self.tables: List[TableData] = []
        self.knowledge_elements: List[KnowledgeElement] = []
        self.document_metadata: Dict[str, Any] = {}
        
    def parse_docx(self, docx_path: str) -> Dict[str, Any]:
        """解析Word文档"""
        if not DOCX_AVAILABLE:
            raise ImportError("需要安装 python-docx: pip install python-docx")
            
        doc = Document(docx_path)
        
        # 提取文档元数据
        self.document_metadata = {
            'source_file': docx_path,
            'parse_time': datetime.now().isoformat(),
            'paragraph_count': len(doc.paragraphs),
            'table_count': len(doc.tables)
        }
        
        # 解析段落（构建章节结构）
        self.sections = self._parse_paragraphs(doc.paragraphs)
        
        # 解析表格
        self.tables = self._parse_tables(doc.tables)
        
        # 提取知识元素
        self.knowledge_elements = self._extract_knowledge()
        
        return self.get_result()
    
    def _parse_paragraphs(self, paragraphs) -> List[DocumentSection]:
        """解析段落，构建章节层级结构"""
        sections = []
        current_section = None
        section_counter = {}
        
        # 标题模式匹配
        heading_patterns = [
            (r'^第[一二三四五六七八九十\d]+章\s*(.+)$', 1),      # 第一章 总则
            (r'^\d+\s*.+$', 2),                                   # 1 总则
            (r'^\d+\.\d+\s*.+$', 3),                              # 1.1 概述
            (r'^\d+\.\d+\.\d+\s*.+$', 4),                         # 1.1.1 总体要求
            (r'^[一二三四五六七八九十]+[、．]\s*(.+)$', 4),        # 一、术语与定义
        ]
        
        for para in paragraphs:
            text = para.text.strip()
            if not text:
                continue
            
            # 检测标题
            for pattern, level in heading_patterns:
                match = re.match(pattern, text)
                if match:
                    # 保存当前章节
                    if current_section:
                        sections.append(current_section)
                    
                    # 创建新章节
                    section_id = f"sec_{level}_{len(sections)}"
                    current_section = DocumentSection(
                        section_id=section_id,
                        title=text,
                        content="",
                        level=level
                    )
                    break
            else:
                # 普通段落，添加到当前章节
                if current_section:
                    current_section.content += text + "\n"
        
        # 添加最后一个章节
        if current_section:
            sections.append(current_section)
        
        return sections
    
    def _parse_tables(self, tables) -> List[TableData]:
        """解析文档中的表格"""
        parsed_tables = []
        
        for idx, table in enumerate(tables):
            table_data = {
                'table_id': f'table_{idx:03d}',
                'headers': [],
                'rows': []
            }
            
            # 提取表头
            if table.rows:
                header_row = table.rows[0]
                table_data['headers'] = [
                    cell.text.strip() for cell in header_row.cells
                ]
            
            # 提取数据行
            for row in table.rows[1:]:
                row_data = [
                    cell.text.strip() for cell in row.cells
                ]
                if any(row_data):  # 跳过空行
                    table_data['rows'].append(row_data)
            
            if table_data['rows']:  # 只保存有数据的表格
                parsed_tables.append(TableData(
                    table_id=table_data['table_id'],
                    headers=table_data['headers'],
                    rows=table_data['rows']
                ))
        
        return parsed_tables
    
    def _extract_knowledge(self) -> List[KnowledgeElement]:
        """从解析结果中提取知识元素"""
        elements = []
        element_counter = {'term': 0, 'rule': 0, 'procedure': 0, 'requirement': 0}
        
        # 从章节中提取
        for section in self.sections:
            # 术语定义（通常在"术语与定义"章节）
            if '术语' in section.title or '定义' in section.title:
                terms = self._extract_terms(section.content)
                for term in terms:
                    element_counter['term'] += 1
                    elements.append(KnowledgeElement(
                        element_id=f"term_{element_counter['term']:04d}",
                        element_type='term',
                        name=term['name'],
                        definition=term['definition'],
                        context=section.section_id,
                        references=[],
                        metadata={'source': '术语章节'}
                    ))
            
            # 规则和要求
            rules = self._extract_rules(section.content)
            for rule in rules:
                element_counter['rule'] += 1
                elements.append(KnowledgeElement(
                    element_id=f"rule_{element_counter['rule']:04d}",
                    element_type=rule['type'],
                    name=rule['name'],
                    definition=rule['content'],
                    context=section.section_id,
                    references=[],
                    metadata={'source': section.title}
                ))
        
        # 从表格中提取
        for table in self.tables:
            table_elements = self._extract_from_table(table)
            elements.extend(table_elements)
        
        return elements
    
    def _extract_terms(self, text: str) -> List[Dict]:
        """从文本中提取术语定义"""
        terms = []
        # 匹配格式：3.1 厂站二次设备：定义内容
        pattern = r'(\d+\.\d+)\s*([^：\n]+)[：:]\s*([^\n]+)'
        matches = re.findall(pattern, text)
        
        for match in matches:
            terms.append({
                'name': match[1].strip(),
                'definition': match[2].strip()
            })
        
        return terms
    
    def _extract_rules(self, text: str) -> List[Dict]:
        """从文本中提取规则"""
        rules = []
        
        # 检测要求类语句
        requirement_patterns = [
            (r'应\s*(.+?)[。.]', 'requirement'),  # 应...
            (r'严禁\s*(.+?)[。.]', 'prohibition'),  # 严禁...
            (r'禁止\s*(.+?)[。.]', 'prohibition'),  # 禁止...
            (r'必须\s*(.+?)[。.]', 'mandatory'),   # 必须...
            (r'宜\s*(.+?)[。.]', 'recommendation'),  # 宜...
        ]
        
        for pattern, rule_type in requirement_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                rules.append({
                    'type': rule_type,
                    'name': f"规则_{len(rules)+1}",
                    'content': match.strip()
                })
        
        # 检测操作动词模式
        action_patterns = [
            r'(投入|退出|打开|连上|断开|合上|拔出|插入|密封|拆除|接入|切换至)\s*([^；;。\n]+)',
        ]
        
        for pattern in action_patterns:
            matches = re.findall(pattern, text)
            for match in matches:
                rules.append({
                    'type': 'procedure',
                    'name': f"流程_{len(rules)+1}",
                    'content': f"{match[0]}{match[1]}"
                })
        
        return rules
    
    def _extract_from_table(self, table: TableData) -> List[KnowledgeElement]:
        """从表格中提取知识元素"""
        elements = []
        
        # 检查是否是措施单格式表格（附录1）
        if len(table.headers) >= 8 and '序号' in table.headers[0]:
            for row in table.rows:
                if len(row) >= 3 and row[0] and row[2]:
                    elements.append(KnowledgeElement(
                        element_id=f"measure_{len(elements)+1}",
                        element_type='procedure',
                        name=row[0],
                        definition=row[2],
                        context=table.table_id,
                        references=[],
                        metadata={'table_id': table.table_id}
                    ))
        
        return elements
    
    def get_result(self) -> Dict[str, Any]:
        """获取解析结果"""
        return {
            'metadata': self.document_metadata,
            'sections': [s.to_dict() for s in self.sections],
            'tables': [t.to_dict() for t in self.tables],
            'knowledge_elements': [e.to_dict() for e in self.knowledge_elements]
        }
    
    def save_results(self, output_path: str, format: str = 'json') -> str:
        """保存解析结果"""
        result = self.get_result()
        
        output_path = Path(output_path)
        output_path.parent.mkdir(parents=True, exist_ok=True)
        
        if format == 'json':
            with open(output_path, 'w', encoding='utf-8') as f:
                json.dump(result, f, ensure_ascii=False, indent=2)
        elif format == 'md':
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write("# 二次作业标准文档解析结果\n\n")
                f.write(f"**解析时间**: {result['metadata']['parse_time']}\n\n")
                f.write(f"**段落数**: {result['metadata']['paragraph_count']}\n\n")
                f.write(f"**表格数**: {result['metadata']['table_count']}\n\n")
                
                # 保存章节
                f.write("## 章节结构\n\n")
                for section in result['sections']:
                    indent = "  " * (section['level'] - 1)
                    f.write(f"{indent}- **{section['title']}**\n")
                
                # 保存表格摘要
                f.write("\n## 表格摘要\n\n")
                for table in result['tables'][:5]:  # 只显示前5个表格
                    f.write(f"\n### {table['table_id']}\n\n")
                    f.write(f"| {' | '.join(table['headers'])} |\n")
                    f.write(f"| {' | '.join(['---'] * len(table['headers']))} |\n")
                    for row in table['rows'][:3]:  # 只显示前3行
                        f.write(f"| {' | '.join(row)} |\n")
        
        return str(output_path)


def main():
    """主函数 - 命令行入口"""
    import argparse
    
    parser = argparse.ArgumentParser(description='电气二次标准文档解析器')
    parser.add_argument('--input', '-i', required=True, help='输入文档路径')
    parser.add_argument('--output', '-o', default='output/parsed.json', help='输出路径')
    parser.add_argument('--format', '-f', choices=['json', 'md'], default='json', help='输出格式')
    
    args = parser.parse_args()
    
    # 创建解析器并运行
    doc_parser = SecondaryDocParser()
    result = doc_parser.parse_docx(args.input)
    output_path = doc_parser.save_results(args.output, args.format)
    
    print(f"✓ 文档解析完成")
    print(f"  - 章节数: {len(result['sections'])}")
    print(f"  - 表格数: {len(result['tables'])}")
    print(f"  - 知识元素数: {len(result['knowledge_elements'])}")
    print(f"  - 输出文件: {output_path}")


if __name__ == '__main__':
    main()
