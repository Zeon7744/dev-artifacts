"""
知识库文件解析：从上传的二进制文件中提取纯文本。

支持格式：
- .txt / .md：UTF-8（容错回退 GBK）解码
- .pdf：pypdf（延迟导入，未安装时抛出明确异常）
- .docx：python-docx（延迟导入，未安装时抛出明确异常）

设计要点：
- 解析库一律在函数内部延迟导入，避免模块顶层 import 失败导致
  整个 RAG 模块不可用；
- 文件大小上限 10MB，超出抛 ValueError；
- 不支持的扩展名抛 ValueError。
"""
import logging
import os
from typing import Optional

logger = logging.getLogger(__name__)

# 支持的扩展名（小写，含点）
SUPPORTED_EXTENSIONS = (".txt", ".md", ".pdf", ".docx")

# 文件大小上限：10MB
MAX_FILE_SIZE = 10 * 1024 * 1024


class ParserDependencyMissing(RuntimeError):
    """解析所需的第三方库未安装（由 API 层映射为 501）。"""

    def __init__(self, package: str, target: str):
        self.package = package
        self.target = target
        super().__init__(f"未安装 {package}，无法解析 {target}")


def _extract_pdf(raw: bytes) -> str:
    """用 pypdf 提取 PDF 文本（延迟导入）。"""
    try:
        from pypdf import PdfReader
    except ImportError:
        raise ParserDependencyMissing("pypdf", "PDF")

    import io

    reader = PdfReader(io.BytesIO(raw))
    parts = []
    for page in reader.pages:
        try:
            text = page.extract_text() or ""
        except Exception as exc:  # 单页解析失败不阻断整体
            logger.warning("PDF 单页解析失败: %s", exc)
            text = ""
        if text.strip():
            parts.append(text.strip())
    return "\n\n".join(parts)


def _extract_docx(raw: bytes) -> str:
    """用 python-docx 提取 Word 文档文本（延迟导入）。"""
    try:
        import docx  # python-docx 的导入名
    except ImportError:
        raise ParserDependencyMissing("python-docx", "DOCX")

    import io

    document = docx.Document(io.BytesIO(raw))
    parts = [p.text for p in document.paragraphs if p.text and p.text.strip()]

    # 表格中的文本也一并提取
    for table in document.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text and cell.text.strip()]
            if cells:
                parts.append(" | ".join(cells))

    return "\n".join(parts)


def _extract_text_file(raw: bytes) -> str:
    """纯文本 / Markdown：优先 UTF-8，失败回退 GBK / latin-1。"""
    for encoding in ("utf-8", "gbk", "latin-1"):
        try:
            return raw.decode(encoding)
        except UnicodeDecodeError:
            continue
    return raw.decode("utf-8", errors="ignore")


def extract_text(filename: str, raw: bytes) -> str:
    """从文件字节中提取纯文本。

    Args:
        filename: 原始文件名（用于判断扩展名与默认标题）
        raw: 文件二进制内容

    Returns:
        提取出的文本内容

    Raises:
        ValueError: 文件为空、超过 10MB、扩展名不支持，或解析后内容为空
        ParserDependencyMissing: 对应解析库未安装（API 层返回 501）
    """
    if not filename:
        raise ValueError("文件名为空")

    ext: Optional[str] = os.path.splitext(filename)[1].lower()
    if ext not in SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"不支持的文件类型: {ext or '无扩展名'}，"
            f"仅支持 {', '.join(SUPPORTED_EXTENSIONS)}"
        )

    if not raw:
        raise ValueError("文件内容为空")
    if len(raw) > MAX_FILE_SIZE:
        raise ValueError(
            f"文件大小 {len(raw)} 字节超过上限 {MAX_FILE_SIZE} 字节（10MB）"
        )

    if ext in (".txt", ".md"):
        text = _extract_text_file(raw)
    elif ext == ".pdf":
        text = _extract_pdf(raw)
    elif ext == ".docx":
        text = _extract_docx(raw)
    else:  # 理论不可达（前面已校验扩展名）
        raise ValueError(f"不支持的文件类型: {ext}")

    if not text or not text.strip():
        raise ValueError(f"文件 {filename} 未解析出有效文本内容")

    return text.strip()
